from docx import Document

document = Document('./test_file/test.docx')  #打开文件demo.docx
for paragraph in document.paragraphs:
    print(paragraph.text)  #打印各段落内容文本

document.add_paragraph(
    'Add new paragraph', style='ListNumber'
)    #添加新段落

document.save('demo.docx') #保存文档